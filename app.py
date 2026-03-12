import streamlit as st
import openpyxl
from openpyxl.styles import Font, Alignment
from openpyxl.cell.cell import MergedCell
import docx
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import calendar
import re
import io
import zipfile
from datetime import date

# ================= 🎨 介面設定 =================
st.set_page_config(page_title="廣慈托嬰中心-行政自動化系統", page_icon="👶", layout="wide")
st.title("📊 托育報表日期自動更新系統")
st.markdown("支援：監測表、冰箱表 (Excel) ｜ 環境檢核表、消防表、教室檢核表 (Word)")

with st.sidebar:
    st.header("📅 全域設定")
    target_year_roc = st.number_input("設定目標民國年份", value=115)
    target_month = st.number_input("設定目標月份", min_value=1, max_value=12, value=3)
    
    st.subheader("🛑 國定假日/停托日")
    holiday_input = st.text_input("輸入日期 (例如: 3/28, 4/4)", help="用逗號隔開數字或月/日。若為雙月表，建議輸入 3/28, 4/4")
    
    st.divider()
    st.info("💡 雙面表單：系統會自動在第一區塊填目標月，第二區塊填次月。")

# ================= 🛠️ 通用邏輯 =================

def get_target_info(year_roc, month, holiday_str):
    year_ad = year_roc + 1911
    _, last_day = calendar.monthrange(year_ad, month)
    holidays_days = []
    
    if holiday_str:
        for item in holiday_str.replace('，', ',').split(','):
            item = item.strip()
            if '/' in item:
                try:
                    m, d = map(int, item.split('/'))
                    if m == month: holidays_days.append(d)
                except: continue
            else:
                try: holidays_days.append(int(item))
                except: continue
    
    workdays = []
    for d in range(1, last_day + 1):
        curr = date(year_ad, month, d)
        if curr.weekday() < 5 and d not in holidays_days:
            workdays.append(curr)
            
    return workdays, holidays_days, last_day

# ================= 📂 Excel 處理邏輯 =================

def process_excel(file_bytes, year_roc, month, workdays):
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes))
    week_map = {0: '一', 1: '二', 2: '三', 3: '四', 4: '五'}
    font_eng = Font(name='Times New Roman', size=12)
    font_chi = Font(name='標楷體', size=12)
    align_center = Alignment(horizontal='center', vertical='center')

    for ws in wb.worksheets:
        for row in ws.iter_rows(min_row=1, max_row=3, min_col=1, max_col=2):
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    cell.value = re.sub(r'\d{2,3}\s*[年./-]\s*\d{1,2}\s*月?', f"{year_roc}年{month:02d}月", cell.value)

        ws_content = "".join([str(cell.value) for row in ws.iter_rows(max_row=10, max_col=5) for cell in row if cell.value])
        
        if "冰箱" in ws_content:
            curr_row = 6
            for d in workdays:
                ws.cell(row=curr_row, column=1).value = f"{d.month}/{d.day}"
                ws.cell(row=curr_row, column=2).value = week_map[d.weekday()]
                curr_row += 2
            while curr_row <= 70:
                if not isinstance(ws.cell(row=curr_row, column=1), MergedCell):
                    ws.cell(row=curr_row, column=1).value = None
                    ws.cell(row=curr_row, column=2).value = None
                curr_row += 1
        elif "監測" in ws_content or "星期" in ws_content:
            date_row = 3
            start_col = 4
            for r in range(1, 6):
                row_vals = [str(ws.cell(row=r, column=c).value) for c in range(1, 6)]
                if any("星期" in v for v in row_vals) or any("日期" in v for v in row_vals):
                    date_row = r
                    break
            for i, d in enumerate(workdays):
                col = start_col + i
                c_day = ws.cell(row=date_row, column=col)
                if not isinstance(c_day, MergedCell):
                    c_day.value = d.day
                    c_day.font = font_eng
                    c_day.alignment = align_center
                c_week = ws.cell(row=date_row+1, column=col)
                if not isinstance(c_week, MergedCell):
                    c_week.value = week_map[d.weekday()]
                    c_week.font = font_chi
                    c_week.alignment = align_center
            for col in range(start_col + len(workdays), 40):
                c_check = ws.cell(row=date_row, column=col)
                if c_check.value and ("餵藥" in str(c_check.value) or "統計" in str(c_check.value)): break
                if not isinstance(c_check, MergedCell): c_check.value = None
                c_next = ws.cell(row=date_row+1, column=col)
                if not isinstance(c_next, MergedCell): c_next.value = None

    out_sim = io.BytesIO()
    wb.save(out_sim)
    return out_sim.getvalue()

# ================= 📄 Word 處理邏輯 =================

def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    for shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(shd)
    new_shd = OxmlElement('w:shd')
    new_shd.set(qn('w:val'), 'clear')
    new_shd.set(qn('w:color'), 'auto')
    new_shd.set(qn('w:fill'), color_hex)
    tcPr.append(new_shd)

def safe_set_word_cell(cell, text, color="FFFFFF"):
    cell.text = ""
    set_cell_shading(cell, color)
    if not text: return
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
    run.font.size = Pt(12)
    paragraph.alignment = 1

def process_docx(file_bytes, target_year_roc, target_month, holiday_input):
    doc = docx.Document(io.BytesIO(file_bytes))
    week_map = {0: '一', 1: '二', 2: '三', 3: '四', 4: '五'}
    
    doc_text = "".join([p.text for p in doc.paragraphs])
    is_fire_safety = "消防" in doc_text or "火源" in doc_text

    # 1. 更新大標題
    for p in doc.paragraphs:
        pattern = r'\d{2,3}\s*年\s*\d{1,2}\s*月'
        if re.search(pattern, p.text):
            original_text = p.text
            new_text = re.sub(pattern, f"{target_year_roc} 年 {target_month:02d} 月", original_text)
            p.text = "" 
            run = p.add_run(new_text)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
            run.font.size = Pt(12)

    # 2. 遍歷表格處理內容
    for table in doc.tables:
        if is_fire_safety:
            # 消防檢查表邏輯 (單月)
            _, holidays, last_day = get_target_info(target_year_roc, target_month, holiday_input)
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    txt = cell.text.strip()
                    if txt.isdigit() and 1 <= int(txt) <= 31:
                        day = int(txt)
                        is_holiday = day in holidays or day > last_day
                        bg_color = "D9D9D9" if is_holiday else "FFFFFF"
                        set_cell_shading(cell, bg_color)
                        if r_idx + 1 < len(table.rows):
                            set_cell_shading(table.rows[r_idx+1].cells[c_idx], bg_color)
        else:
            # 環境/教室檢核表 (支援雙月份填充)
            date_rows_found = []
            for i, row in enumerate(table.rows):
                combined_row_text = "".join([c.text for c in row.cells])
                if "日期" in combined_row_text or "項目" in combined_row_text:
                    date_rows_found.append(i)
            
            # 針對找到的每一個日期區塊進行處理
            for block_idx, row_idx in enumerate(date_rows_found):
                calc_month = target_month + block_idx
                calc_year = target_year_roc
                if calc_month > 12:
                    calc_month -= 12
                    calc_year += 1
                
                # 重新計算該月份的工作日
                block_workdays, _, _ = get_target_info(calc_year, calc_month, holiday_input)
                
                date_row = table.rows[row_idx]
                weekday_row = table.rows[row_idx + 1] if row_idx + 1 < len(table.rows) else None
                
                start_col = 0
                for c_idx, cell in enumerate(date_row.cells):
                    if cell.text.strip().isdigit():
                        start_col = c_idx
                        break
                
                for i in range(start_col, len(date_row.cells)):
                    idx = i - start_col
                    if idx < len(block_workdays):
                        d = block_workdays[idx]
                        safe_set_word_cell(date_row.cells[i], str(d.day))
                        if weekday_row: 
                            safe_set_word_cell(weekday_row.cells[i], week_map[d.weekday()])
                    else:
                        # 清空多餘格子
                        safe_set_word_cell(date_row.cells[i], "")
                        if weekday_row: 
                            safe_set_word_cell(weekday_row.cells[i], "")

    out_sim = io.BytesIO()
    doc.save(out_sim)
    return out_sim.getvalue()

# ================= 🚀 執行介面 =================

uploaded_files = st.file_uploader("📂 上傳報表 (Excel 或 Word)", type=["xlsx", "docx"], accept_multiple_files=True)

if uploaded_files:
    if st.button("🚀 開始批次更新"):
        # Excel 預設抓主月份
        main_workdays, _, _ = get_target_info(target_year_roc, target_month, holiday_input)
        
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w") as zf:
            for uploaded_file in uploaded_files:
                fname = uploaded_file.name
                f_bytes = uploaded_file.read()
                try:
                    if fname.endswith(".xlsx"):
                        processed_data = process_excel(f_bytes, target_year_roc, target_month, main_workdays)
                    elif fname.endswith(".docx"):
                        # Word 調用支援雙月的 process_docx
                        processed_data = process_docx(f_bytes, target_year_roc, target_month, holiday_input)
                    
                    new_name = f"更新_{target_year_roc}年{target_month}月_{fname}"
                    zf.writestr(new_name, processed_data)
                    st.write(f"✅ 已完成: {fname}")
                except Exception as e:
                    st.error(f"❌ 處理 {fname} 出錯: {e}")
        
        st.success(f"🎉 全部處理完成！")
        st.download_button(
            label="📥 下載更新後的報表 (ZIP)",
            data=zip_buffer.getvalue(),
            file_name=f"廣慈報表_{target_year_roc}年{target_month}月.zip",
            mime="application/zip"
        )
